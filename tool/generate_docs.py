"""
Generate Word (.docx) versions of the project documentation.

Run once after installing python-docx:

    pip install python-docx
    python tool/generate_docs.py

Outputs:
    docs/PROJECT_REPORT.docx
    docs/USER_GUIDE.docx
    docs/INSTALLATION.docx
    docs/README.docx

The markdown files in docs/ and README.md remain unchanged.

This is intentionally a pure Python script (no extra CLI tools such as
pandoc required) so that any evaluator with Python 3.8+ can regenerate the
Word documents.
"""

from __future__ import annotations

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "Missing dependency: python-docx.\n"
        "Install it with:\n\n    pip install python-docx\n\n"
    )
    sys.exit(1)


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"

PRIMARY = RGBColor(0x1E, 0x88, 0xE5)
MUTED = RGBColor(0x55, 0x55, 0x55)


def _set_cell_shading(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = PRIMARY if level <= 2 else RGBColor(0, 0, 0)
    sizes = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}
    run.font.size = Pt(sizes.get(level, 11))
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


INLINE_RE = re.compile(r"(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\))")


def _add_inline(paragraph, text: str) -> None:
    """Add a text line that may contain **bold**, *italic*, `code`, [links]."""
    idx = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > idx:
            paragraph.add_run(text[idx : m.start()])
        if m.group(2) is not None:  # **bold**
            r = paragraph.add_run(m.group(2))
            r.bold = True
        elif m.group(3) is not None:  # *italic*
            r = paragraph.add_run(m.group(3))
            r.italic = True
        elif m.group(4) is not None:  # `code`
            r = paragraph.add_run(m.group(4))
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = PRIMARY
        elif m.group(5) is not None:  # [label](url)
            r = paragraph.add_run(m.group(5))
            r.font.color.rgb = PRIMARY
            r.underline = True
        idx = m.end()
    if idx < len(text):
        paragraph.add_run(text[idx:])


def _add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # light grey shading on the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F3F5")
    pPr.append(shd)


def _parse_table(lines: list[str]) -> list[list[str]] | None:
    """Parse a simple GFM pipe table. Returns rows of cells or None."""
    if len(lines) < 2 or not lines[1].strip().startswith("|"):
        return None
    rows = [
        [c.strip() for c in ln.strip().strip("|").split("|")]
        for ln in lines
        if ln.strip().startswith("|")
    ]
    if len(rows) < 2:
        return None
    # drop the --- separator row
    sep = rows[1]
    if all(re.match(r":?-+:?", c or "") for c in sep):
        rows.pop(1)
    return rows


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=cols)
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = tbl.rows[i].cells[j]
            text = row[j] if j < len(row) else ""
            cell.text = ""
            _add_inline(cell.paragraphs[0], text)
            if i == 0:
                _set_cell_shading(cell, "1E88E5")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def convert(md_path: Path, out_path: Path, title: str) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    # Title block
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(title)
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("City Guide Mobile Application - eProject submission")
    sr.italic = True
    sr.font.color.rgb = MUTED

    doc.add_paragraph()  # spacer

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        # code fence
        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < n and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            _add_code_block(doc, block)
            continue

        # tables
        if stripped.startswith("|") and i + 1 < n and lines[i + 1].strip().startswith("|"):
            table_lines: list[str] = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table(table_lines)
            if rows:
                _add_table(doc, rows)
                continue
            # fallthrough: treat as paragraphs
            for tl in table_lines:
                _add_inline(doc.add_paragraph(), tl)
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            _add_heading(doc, m.group(2).strip(), len(m.group(1)))
            i += 1
            continue

        # horizontal rule
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # bulleted list
        if re.match(r"^\s*[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, re.sub(r"^\s*[-*]\s+", "", stripped))
            i += 1
            continue

        # numbered list
        if re.match(r"^\s*\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\s*\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped == "":
            doc.add_paragraph()
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run(stripped.lstrip("> ").rstrip())
            run.italic = True
            run.font.color.rgb = MUTED
            i += 1
            continue

        # plain paragraph
        _add_inline(doc.add_paragraph(), stripped)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    rel = out_path.relative_to(ROOT)
    print(f"  wrote {rel}  ({out_path.stat().st_size // 1024} KB)")


def main() -> None:
    pairs = [
        (ROOT / "README.md", DOCS / "README.docx", "City Guide - README"),
        (DOCS / "PROJECT_REPORT.md", DOCS / "PROJECT_REPORT.docx", "City Guide - Project Report"),
        (DOCS / "INSTALLATION.md", DOCS / "INSTALLATION.docx", "City Guide - Installation Guide"),
        (DOCS / "USER_GUIDE.md", DOCS / "USER_GUIDE.docx", "City Guide - User Guide"),
    ]
    print("Generating Word documents...")
    for md, out, title in pairs:
        if not md.exists():
            print(f"  skip (missing): {md.relative_to(ROOT)}")
            continue
        convert(md, out, title)
    print("Done.")


if __name__ == "__main__":
    main()
